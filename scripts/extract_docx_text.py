"""
DOCX Text Extraction using python-docx (pure Python)
Optimized for Vercel serverless execution - fast and memory-efficient.
"""

import sys
import json
from typing import Dict
from docx import Document


def extract_text_from_docx(docx_path: str) -> Dict[str, any]:
    """
    Extract text from a DOCX file using python-docx.
    Optimized for speed - efficient text extraction.
    
    Args:
        docx_path: Path to the DOCX file
    
    Returns:
        Dictionary containing extracted text and metadata
    """
    try:
        # Load the document
        doc = Document(docx_path)
        
        # Extract text from all paragraphs (fast)
        text_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        text_parts.append(text)
        
        # Join all text parts
        full_text = '\n'.join(text_parts)
        cleaned_text = full_text.strip()
        
        # Get basic metadata (fast)
        core_props = doc.core_properties
        
        return {
            "success": True,
            "text": cleaned_text,
            "character_count": len(cleaned_text),
            "word_count": len(cleaned_text.split()),
            "metadata": {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "extractor": "python-docx"
            }
        }
    
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }


def main():
    """Main function for CLI usage."""
    if len(sys.argv) < 2:
        print(json.dumps({
            "success": False,
            "error": "No DOCX file path provided"
        }))
        sys.exit(1)
    
    docx_path = sys.argv[1]
    result = extract_text_from_docx(docx_path)
    print(json.dumps(result, indent=2))
    
    if not result["success"]:
        sys.exit(1)


if __name__ == "__main__":
    main()
